from docx import Document

# Création du document
doc = Document()

# Ajout du titre principal
doc.add_heading('Révision Algorithme : Bases, Boucles et Fonctions', level=1)

# Ajout du sommaire
doc.add_heading('Sommaire', level=2)
sommaire_sections = [
    "Introduction à l’algorithme",
    "Les Variables et Types de Données",
    "Les Structures Conditionnelles (Tests)",
    "Les Boucles",
    "Les Fonctions et Procédures",
    "Exercices et Applications",
    "Conclusion et Conseils pour l’Examen"
]

for i, section in enumerate(sommaire_sections, 1):
    doc.add_paragraph(f"{i}. {section}", style="List Number")

# Ajout des sections détaillées
sections_content = {
    "Introduction à l’algorithme": [
        "Un algorithme est une suite d'instructions permettant de résoudre un problème de manière ordonnée.",
        "Il comprend trois parties principales :",
        "- Les entrées : les données à traiter.",
        "- Le traitement : les opérations effectuées sur les entrées.",
        "- Les sorties : les résultats obtenus.",
        "Exemple de structure générale d’un algorithme :",
        "Début\n    Instructions\nFin"
    ],
    "Les Variables et Types de Données": [
        "Une variable est un espace mémoire permettant de stocker une valeur.",
        "Types de données courants :",
        "- Entier : nombres sans décimale.",
        "- Réel : nombres avec décimale.",
        "- Chaîne : texte ou suite de caractères.",
        "- Booléen : VRAI ou FAUX.",
        "Exemple de déclaration en pseudo-code :",
        "Var age : Entier\nVar nom : Chaîne"
    ],
    "Les Structures Conditionnelles (Tests)": [
        "SI...ALORS...SINON :",
        "Si condition Alors\n    Instructions\nSinon\n    Autres instructions\nFinSi",
        "Exemple :",
        "Si age >= 18 Alors\n    Afficher 'Majeur'\nSinon\n    Afficher 'Mineur'\nFinSi"
    ],
    "Les Boucles": [
        "Boucle TANT QUE :",
        "TantQue condition Faire\n    Instructions\nFinTantQue",
        "Boucle POUR :",
        "Pour i de 1 à N Faire\n    Instructions\nFinPour"
    ],
    "Les Fonctions et Procédures": [
        "Une fonction est un bloc d'instructions qui retourne une valeur.",
        "Exemple de déclaration :",
        "Fonction Carré(x : Entier) : Entier\nDébut\n    Retourner x * x\nFinFonction",
        "Appel d'une fonction :",
        "Résultat ← Carré(5) // Résultat = 25"
    ],
    "Exercices et Applications": [
        "1. Algorithmes à compléter",
        "2. Exercices sur les boucles",
        "3. Exercices sur les conditions",
        "4. Mise en pratique avec des fonctions"
    ],
    "Conclusion et Conseils pour l’Examen": [
        "- Comprendre les bases, les tests et les boucles.",
        "- Décomposer le problème avant d’écrire un algorithme.",
        "- Tester chaque partie avant d’exécuter l’ensemble."
    ]
}

for section, content in sections_content.items():
    doc.add_heading(section, level=2)
    for paragraph in content:
        doc.add_paragraph(paragraph)

# Sauvegarde du fichier
file_path = "/Users/pyinv/Dev/automat/Revision_Algorithme.docx"
doc.save(file_path)

file_path
